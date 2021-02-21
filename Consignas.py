import psycopg2
import xlrd
import docx

from urllib.request import urlretrieve
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from Creacion_db import *




# No se hará un ya que esta creación es específica.. si funciones para repetirlas

#El resutlado no se sube a DB, solo consulta. Genera texto para usar en WORD y enviar al alumno

conn = psycopg2.connect(
    host = 'localhost',
    database = 'instanciaMarzo',
    port = 5432,
    user = 'postgres',
    password = 'postgres'
)


conn.autocommit = True  

#Cursor
cursor = conn.cursor()


def consignas(tabla_curso, tabla_datos_ejercicio):



    #Creamos un objeto para usarlo- Document sera el punto de entrada
    doc = docx.Document()


    my_image = doc.add_picture('logo_escuela.png')
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    

    texto_escuela = doc.add_paragraph()
    texto_escuela.add_run("E.E.N° 178 'San José de Calasanz'")
    texto_escuela.runs[0].font.size = Pt(8)
    texto_escuela.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
    texto_escuela2 = doc.add_paragraph()
    texto_escuela2.add_run("General Vedia. Chaco.")
    texto_escuela2.runs[0].font.size = Pt(8)
    texto_escuela2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    salto_de_linea_logo = doc.add_paragraph()



    #Linea doble encabezado
    linea_superior1 = doc.add_paragraph() 
    linea_superior1.add_run("▬" *37)
    linea_superior1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    linea_superior2 = doc.add_paragraph() 
    linea_superior2.add_run("▬" *37)
    linea_superior2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #Interior encabezado
    encabezado = doc.add_heading(f"--- CONSIGNAS Y DATOS POR NOMBRE DE ESTUDIANTE ---", level=2)
    encabezado.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    encabezado2 = doc.add_heading("".join(tabla_curso.split("real_")), level=2)
    encabezado2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    encabezado2.runs[0].font.color.rgb = RGBColor(180, 0, 0)
    #Los encabezados también son runs



    #Aclaración para estudiante
    aclaracion = doc.add_paragraph()
    aclaracion.add_run("Solo debés utilizar las consignas y datos que corresponden con tu nombre :)")
    aclaracion.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    linea_inferior = doc.add_paragraph() 
    linea_inferior.add_run("▬" *37)
    linea_inferior.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    salto_linea_portada = doc.add_paragraph()
    salto_linea_portada2 = doc.add_paragraph()
    salto_linea_portada3 = doc.add_paragraph()

    profesor = doc.add_paragraph()
    profesor.add_run("Profesor").bold = True
    profesor.add_run(": Wilson Díaz")
    profesor.runs[0].font.size = Pt(12)
    profesor.runs[1].font.size = Pt(12)
    formato_profesor = profesor.paragraph_format
    formato_profesor.first_line_indent = Cm(1)

    salto_linea_portada4 = doc.add_paragraph()

    asignatura = doc.add_paragraph()
    asignatura.add_run("Asignatura").bold = True
    asignatura.add_run(": Educación Tecnológica")
    asignatura.runs[0].font.size = Pt(12)
    asignatura.runs[1].font.size = Pt(12)
    formato_asignatura = asignatura.paragraph_format
    formato_asignatura.first_line_indent = Cm(1)

    salto_linea_portada4 = doc.add_paragraph()

    año = doc.add_paragraph()
    año.add_run("Fecha").bold = True
    año.add_run(": Marzo/2021")
    año.runs[0].font.size = Pt(12)
    año.runs[1].font.size = Pt(12)
    formato_año = año.paragraph_format
    formato_año.first_line_indent = Cm(1)

    salto_linea_portada5 = doc.add_paragraph()

    curso_div = doc.add_paragraph()
    curso_div.add_run("3er año").bold = True
    curso_div.runs[0].font.size = Pt(12)
    formato_curso_div = curso_div.paragraph_format
    formato_curso_div.first_line_indent = Cm(1)






    salto_de_pagina = doc.add_page_break() #El salto de página

     

    


    #Aquí se sabe que la tabla del curso su segundo atributo es "estudiante" (ajustar si usas otra tabla)

    #Conocer cantidad de registros(estudiantes) que hay en la tabla elegida
    cursor.execute(f"SELECT COUNT(*) FROM {tabla_curso}")
    estudiantes = cursor.fetchall()[0][0] #Los [0][0] es para eliminar su tupla dentro de una lista así da solo el int
    

    



    for cada_estudiante in range(estudiantes+1): #+1 para eliminar conteo desde 0
        

        
                        #Se conoce q estudiante es donde esta el nombre dentro de la tabla   
        cursor.execute(f"SELECT estudiante FROM {tabla_curso} WHERE id={cada_estudiante}") 
        alumno = cursor.fetchall()

        for i in range(len(alumno)): #No salía con alumno[0][0] así que tuve que hacerlo así
            alumno = alumno[i][0]


            #___________________________________________________________________________________________________________



            estudiante = doc.add_paragraph(style = "List Bullet")
            estudiante.add_run("Estudiante").bold = True 
            estudiante.runs[0].font.size = Pt(12)
            estudiante.runs[0].underline = True
            estudiante.runs[0].font.color.rgb = RGBColor(180, 0, 0)
            formato_estudiante = estudiante.paragraph_format #Creando el objeto formato para manejarlo
            formato_estudiante.first_line_indent = Cm(0.5) #Un cm de sangria

            estudiante.add_run (f":  {str(alumno).upper()}").bold= True
            estudiante.runs[1].font.size = Pt(12)
            estudiante.runs[1].font.color.rgb = RGBColor(180, 0, 0)
            
            espacio_mas = doc.add_paragraph()

            #                                ------------ EJERCICIOS -------------

            #Generando punto 9) a)                                             #id en este caso es igual al id foraneo
            cursor.execute(f"Select ej_9_a FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()
            
            linea_libre2 = doc.add_paragraph()
            
            consigna = doc.add_paragraph()
            consigna.add_run("Consigna y/o datos del punto 9)a) :").bold = True
            consigna_formato = consigna.paragraph_format
            consigna_formato.first_line_indent = Cm(0.5)
            


            cuerpo_consigna = doc.add_paragraph()
            cuerpo_consigna.add_run("¿Qué ")
            cuerpo_consigna.add_run("grosor ").bold = True
            cuerpo_consigna.add_run("de cable debés utilizar si el consumo de tu estufa es de ")
            cuerpo_consigna.add_run(f"{dato[0][0]} amperes?").bold = True

            #___________________________________________________________________________________________________________


            
            #Generando punto 9) B)
            cursor.execute(f"Select ej_9_b FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()
            
            linea_libre3 = doc.add_paragraph()

            consigna2 = doc.add_paragraph() 
            consigna2.add_run("Consigna y/o datos del punto 9)b) :").bold = True
            consigna_formato2 = consigna2.paragraph_format
            consigna_formato2.first_line_indent = Cm(0.5)


            cuerpo_consigna2 = doc.add_paragraph()
            cuerpo_consigna2.add_run("¿Si el cable que compraste es de ")
            cuerpo_consigna2.add_run(f"{dato[0][0]} mm").bold = True
            cuerpo_consigna2.add_run(", cuántos ")
            cuerpo_consigna2.add_run("amperes").bold = True
            cuerpo_consigna2.add_run(" soportará de consumo antes que comience a calentarse?")
            cuerpo_consigna2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

  
            #___________________________________________________________________________________________________________


            
            #Generando punto 9) C)

            linea_libre4 = doc.add_paragraph()

            consigna3 = doc.add_paragraph() 
            consigna3.add_run("Consigna y/o datos del punto 9)c) :").bold = True
            consigna_formato3 = consigna3.paragraph_format
            consigna_formato3.first_line_indent = Cm(0.5)

            
            cuerpo_consigna2 = doc.add_paragraph()
            cuerpo_consigna2.add_run("Si enchufamos EN UNA MISMA zapatilla los siguientes electrodomésticos:")



            cursor.execute(f"Select ej_9_c1 FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()

            cuerpo_consigna3 = doc.add_paragraph()
            cuerpo_consigna3.add_run(" - Nuestra PC Gamer que consume ")
            cuerpo_consigna3.add_run(f"{dato[0][0][0]} amperes.").bold = True


            cursor.execute(f"Select ej_9_c2 FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato1 = cursor.fetchall()

            cuerpo_consigna4 = doc.add_paragraph()
            cuerpo_consigna4.add_run(" - Una heladera con congelador que consume ")
            cuerpo_consigna4.add_run(f"{dato1[0][0]} amperes.").bold = True


            cursor.execute(f"Select ej_9_c3 FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato2 = cursor.fetchall()

            cuerpo_consigna5 = doc.add_paragraph()
            cuerpo_consigna5.add_run(" - Un minicomponente que consume ")
            cuerpo_consigna5.add_run(f"{dato2[0][0]} amperes.").bold = True

            cuerpo_consigna6 = doc.add_paragraph()
            cuerpo_consigna6.add_run("PREGUNTA: ¿Qué ")
            cuerpo_consigna6.add_run("grosor").bold = True
            cuerpo_consigna6.add_run(" debería tener el cable que alimenta a esa zapatilla?")

            


            #___________________________________________________________________________________________________________

                        
            #Generando punto 11)
            cursor.execute(f"Select ej_11 FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()

            linea_libre4 = doc.add_paragraph()
            consigna4 = doc.add_paragraph() 
            consigna4.add_run("Consigna y/o datos del punto 11) :").bold = True
            consigna_formato4 = consigna4.paragraph_format
            consigna_formato4.first_line_indent = Cm(0.5)

            cuerpo_consigna7 = doc.add_paragraph()
            cuerpo_consigna7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            cuerpo_consigna7.add_run ("Necesito un alargue para alimentar un acondicionador de aire que funciona a una potencia de ")
            cuerpo_consigna7.add_run(f"{dato[0][0]} watts.").bold = True
            



            cuerpo_consigna8 = doc.add_paragraph()
            cuerpo_consigna8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            cuerpo_consigna8.add_run("PREGUNTA: ¿Cuántos ")
            cuerpo_consigna8.add_run(f"milímetros").bold = True
            cuerpo_consigna8.add_run(" debe tener la sección de cable (grosor) del alargue que debemos usar para alimntar al acondicionador de aire?")
            
            cuerpo_consigna8b = doc.add_paragraph("-dato: el voltaje en Argentina siempre es de 220 V.-")
            




            #___________________________________________________________________________________________________________



            
            #Generando punto 13) #Use a y b en el mismo ejercicio para hacerlo más rápido
            cursor.execute(f"Select ej_13_a FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()
            cursor.execute(f"Select ej_13_b FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato1 = cursor.fetchall()

            linea_libre5 = doc.add_paragraph()

            consigna5 = doc.add_paragraph()
            consigna5.add_run("Consigna y/o datos del punto 13) :").bold = True
            consigna_formato5 = consigna5.paragraph_format
            consigna_formato5.first_line_indent = Cm(0.5)

            cuerpo_consigna9 = doc.add_paragraph()
            cuerpo_consigna9.add_run("Pasar a decimal ---> ")
            cuerpo_consigna9.add_run(f" {dato[0][0]}.").bold = True

            cuerpo_consigna10 = doc.add_paragraph()
            cuerpo_consigna10.add_run("Pasar a decimal ---> ")
            cuerpo_consigna10.add_run(f" {dato1[0][0]}.").bold = True

           

            #___________________________________________________________________________________________________________




            
            #Generando punto 14)
            cursor.execute(f"Select ej_14 FROM {tabla_datos_ejercicio} WHERE id={cada_estudiante}") 
            dato = cursor.fetchall()


            linea_libre6 = doc.add_paragraph()

            consigna6 = doc.add_paragraph()
            consigna6.add_run("Consigna y/o datos del punto 14) :").bold = True
            consigna_formato6 = consigna6.paragraph_format
            consigna_formato6.first_line_indent = Cm(0.5)

            cuerpo_consigna11 = doc.add_paragraph()
            cuerpo_consigna11.add_run("Calcular cuál es la capacidad real en ")
            cuerpo_consigna11.add_run("GIBIBYTES").bold = True
            cuerpo_consigna11.add_run(" que tiene una unidad de almacenamiento que compraste (ya sea pendrive o disco rígido, da igual), si la etiqueta dice que es de ")
            cuerpo_consigna11.add_run(f"{dato[0][0]} Gigas.").bold = True
            cuerpo_consigna11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            todos_los_parrafos = doc.paragraphs
            for cada_parrafo in todos_los_parrafos:
                cada_parrafo.paragraph_format.space_after = Pt(3) #A todos los párrafos le daremos menos interlineado

            ultimo_espacio = doc.add_paragraph()
            ultimo_espacio2 = doc.add_paragraph()

            linea_final = doc.add_paragraph() 
            linea_final.add_run("_" *100)
            linea_final.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            if cada_estudiante != estudiantes:
                ultimo_salto_de_pagina = doc.add_page_break()



        
        #Aquí le pondrás
        doc.save("Ejercicios previa tecno marzo2021.docx")

    print("Proceso terminado con éxito.")
        
        




#Aquí vas a especificar las bases de datos que se utilizará
#consignas('real_previas_marzo2021_tecnologia', 'ejer_previas_marzo2021_tecno')


